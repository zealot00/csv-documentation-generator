"""Word document generator for CSV documentation - Professional GMP Style"""

import re
from typing import Dict, Any, Optional, List
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell


# =============================================================================
# Professional Color Scheme - GMP/Pharma Industry Style
# =============================================================================
class Colors:
    """Professional color palette for GMP documentation"""

    # Primary colors - Professional blue
    PRIMARY_DARK = RGBColor(30, 58, 95)  # #1E3A5F - Deep blue
    PRIMARY_MEDIUM = RGBColor(37, 99, 235)  # #2563EB - Royal blue (better contrast)
    PRIMARY_LIGHT = RGBColor(59, 130, 246)  # Blue for links

    # Text colors
    TEXT_PRIMARY = RGBColor(30, 41, 59)  # #1E293B - Dark slate (high contrast)
    TEXT_SECONDARY = RGBColor(71, 85, 105)  # #475569 - Slate gray
    TEXT_WHITE = RGBColor(255, 255, 255)  # White

    # Table colors - Professional with good contrast
    TABLE_HEADER_BG = RGBColor(37, 99, 235)  # #2563EB - Royal blue
    TABLE_HEADER_TEXT = RGBColor(255, 255, 255)  # White
    TABLE_ROW_EVEN = RGBColor(255, 255, 255)  # White
    TABLE_ROW_ODD = RGBColor(241, 245, 249)  # #F1F5F9 - Light slate
    TABLE_BORDER = RGBColor(203, 213, 225)  # #CBD5E1 - Slate border

    # Status colors
    STATUS_PASS = RGBColor(5, 150, 105)  # #059669 - Green
    STATUS_FAIL = RGBColor(185, 28, 28)  # #B91C1C - Red
    STATUS_PENDING = RGBColor(180, 83, 9)  # #B45309 - Amber
    STATUS_NA = RGBColor(107, 114, 128)  # #6B7280 - Gray

    # Priority colors
    PRIORITY_MUST = RGBColor(185, 28, 28)  # #B91C1C - Red
    PRIORITY_SHOULD = RGBColor(180, 83, 9)  # #B45309 - Amber
    PRIORITY_COULD = RGBColor(5, 150, 105)  # #059669 - Green

    # Font settings
    FONT_CN = "宋体"  # Chinese font
    FONT_EN = "Arial"  # English font


# =============================================================================
# Document Type Info
# =============================================================================
DOCUMENT_INFO = {
    "vp": {"name_cn": "验证计划", "name_en": "Validation Plan", "abbr": "VP"},
    "urs": {
        "name_cn": "用户需求规格",
        "name_en": "User Requirements Specification",
        "abbr": "URS",
    },
    "fs": {"name_cn": "功能规格", "name_en": "Functional Specification", "abbr": "FS"},
    "ts": {"name_cn": "技术规格", "name_en": "Technical Specification", "abbr": "TS"},
    "ra": {"name_cn": "风险评估", "name_en": "Risk Assessment", "abbr": "RA"},
    "iq": {
        "name_cn": "安装确认",
        "name_en": "Installation Qualification",
        "abbr": "IQ",
    },
    "oq": {"name_cn": "操作确认", "name_en": "Operational Qualification", "abbr": "OQ"},
    "pq": {"name_cn": "性能确认", "name_en": "Performance Qualification", "abbr": "PQ"},
    "vsr": {
        "name_cn": "验证总结报告",
        "name_en": "Validation Summary Report",
        "abbr": "VSR",
    },
}


# =============================================================================
# Word Generator Class
# =============================================================================
class WordGenerator:
    """Generate Word documents with professional GMP styling"""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.doc_info = None

    def generate(
        self,
        doc_type: str,
        content: str,
        variables: Dict[str, str],
        filename: Optional[str] = None,
    ) -> str:
        """Generate Word document with professional styling"""

        doc = Document()
        self.doc_info = DOCUMENT_INFO.get(
            doc_type,
            {"name_cn": doc_type, "name_en": doc_type, "abbr": doc_type.upper()},
        )

        # Process variables
        for key, value in variables.items():
            content = content.replace(f"{{{key}}}", str(value))

        # Setup document styles
        self._setup_document_styles(doc)

        # Check if content starts with title (封面标记)
        lines = content.split("\n")

        # Generate cover page first
        cover_data = self._extract_cover_info(lines, variables)
        self._generate_cover_page(doc, cover_data)

        # Add page break after cover
        doc.add_page_break()

        # Process remaining content
        remaining_lines = []
        in_cover = False
        for line in lines:
            stripped = line.strip()
            # Skip cover page content markers
            if stripped.startswith("# ") and (
                "用户需求" in stripped
                or "验证计划" in stripped
                or "功能规格" in stripped
                or "技术规格" in stripped
                or "风险评估" in stripped
                or "安装确认" in stripped
                or "操作确认" in stripped
                or "性能确认" in stripped
                or "Validation Plan" in stripped
                or "User Requirements" in stripped
                or "Functional Specification" in stripped
                or "Technical Specification" in stripped
                or "Risk Assessment" in stripped
                or "Qualification" in stripped
                or "Summary" in stripped
            ):
                in_cover = True
                continue
            if in_cover and stripped == "---":
                in_cover = False
                continue
            if not in_cover:
                remaining_lines.append(line)

        content = "\n".join(remaining_lines)

        # Setup headers and footers
        self._setup_headers_footers(doc, variables)

        # Process content
        self._process_content(doc, content)

        # Save document
        if filename is None:
            filename = f"{doc_type.upper()}_{self._sanitize_filename(variables.get('PROJECT_NAME', 'Document'))}.docx"

        output_path = self.output_dir / filename
        doc.save(str(output_path))

        return str(output_path)

    def _setup_document_styles(self, doc: Document):
        """Setup document-wide styles with unified fonts"""
        # Normal style
        style = doc.styles["Normal"]
        font = style.font
        font.name = Colors.FONT_EN
        font.size = Pt(11)
        font.color.rgb = Colors.TEXT_PRIMARY
        # Set Chinese font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

        # Heading styles
        for level in range(1, 4):
            if f"Heading {level}" in doc.styles:
                heading_style = doc.styles[f"Heading {level}"]
                heading_font = heading_style.font
                heading_font.name = Colors.FONT_EN
                heading_font.bold = True
                heading_font.color.rgb = Colors.PRIMARY_DARK
                heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)
                if level == 1:
                    heading_font.size = Pt(16)
                elif level == 2:
                    heading_font.size = Pt(14)
                else:
                    heading_font.size = Pt(12)

    def _extract_cover_info(
        self, lines: List[str], variables: Dict[str, str]
    ) -> Dict[str, str]:
        """Extract information for cover page"""
        return {
            "doc_name_cn": self.doc_info.get("name_cn", ""),
            "doc_name_en": self.doc_info.get("name_en", ""),
            "doc_abbr": self.doc_info.get("abbr", ""),
            "project": variables.get("PROJECT_NAME", "[项目名称]"),
            "system": variables.get("SYSTEM_NAME", "[系统名称]"),
            "version": variables.get("VERSION", "1.0"),
            "date": variables.get("DATE", "[日期]"),
            "doc_id": variables.get("DOC_ID", "[文档编号]"),
            "author": variables.get("AUTHOR", "[作者]"),
            "reviewer": variables.get("REVIEWER", "[审核人]"),
            "approver": variables.get("APPROVER", "[批准人]"),
            "category": variables.get("GAMP_CATEGORY", ""),
        }

    def _generate_cover_page(self, doc: Document, data: Dict[str, str]):
        """Generate professional cover page"""
        # Title section - centered, large
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"{data['doc_name_cn']}\n{data['doc_name_en']}")
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = Colors.PRIMARY_DARK
        title_run.font.name = Colors.FONT_EN
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

        # Add spacing
        doc.add_paragraph()

        # Project & System info table
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = "Table Grid"
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        for cell in info_table.columns[0].cells:
            cell.width = Cm(5)
        for cell in info_table.columns[1].cells:
            cell.width = Cm(12)

        # Fill info
        self._fill_cover_row(
            info_table.rows[0].cells[0], "项目 / Project:", data["project"]
        )
        self._fill_cover_row(
            info_table.rows[0].cells[1], "系统 / System:", data["system"]
        )
        self._fill_cover_row(
            info_table.rows[1].cells[0], "文档编号 / Doc ID:", data["doc_id"]
        )
        self._fill_cover_row(
            info_table.rows[1].cells[1], "版本 / Version:", data["version"]
        )
        self._fill_cover_row(
            info_table.rows[2].cells[0],
            "GAMP 分类 / Category:",
            f"Category {data['category']}",
        )
        self._fill_cover_row(info_table.rows[2].cells[1], "日期 / Date:", data["date"])

        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Approval section title
        approval_title = doc.add_paragraph()
        approval_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        approval_run = approval_title.add_run("审批签字表 / Approval Signature")
        approval_run.font.size = Pt(14)
        approval_run.font.bold = True
        approval_run.font.color.rgb = Colors.PRIMARY_DARK
        approval_run.font.name = Colors.FONT_EN
        approval_run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

        # Approval table
        approval_table = doc.add_table(rows=4, cols=4)
        approval_table.style = "Table Grid"

        # Set header row
        headers = ["角色 / Role", "签名 / Signature", "日期 / Date", "备注 / Notes"]
        header_cells = approval_table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            self._set_header_cell_text(cell, header)

        # Fill approval rows
        roles = [
            ("作者 / Author", data["author"]),
            ("审核 / Reviewer", data["reviewer"]),
            ("批准 / Approver", data["approver"]),
        ]

        for idx, (role, name) in enumerate(roles):
            row = approval_table.rows[idx + 1].cells
            row[0].text = role
            row[1].text = name
            row[2].text = ""
            row[3].text = ""
            self._set_data_cell_text(row[0], role)
            self._set_data_cell_text(row[1], name)
            self._set_data_cell_text(row[2], "")
            self._set_data_cell_text(row[3], "")

    def _fill_cover_row(self, cell, label, value):
        """Fill a cover info row"""
        cell.text = f"{label}\n{value}"
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)

        # Style runs
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = Colors.TEXT_PRIMARY
            run.font.name = Colors.FONT_EN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

    def _set_header_cell_text(self, cell: _Cell, text: str):
        """Set header cell text and styling - FIXED version"""
        # Set background color first
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "2563EB")  # Royal blue
        cell._tc.get_or_add_tcPr().append(shading)

        # Clear and rebuild paragraph with styled run
        para = cell.paragraphs[0]
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add styled run with explicit text
        run = para.add_run(text)
        run.font.bold = True
        run.font.color.rgb = Colors.TEXT_WHITE
        run.font.size = Pt(10)
        run.font.name = Colors.FONT_EN
        run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

    def _set_data_cell_text(self, cell: _Cell, text: str):
        """Set data cell text and styling"""
        cell.text = text
        para = cell.paragraphs[0]
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = Colors.TEXT_PRIMARY
            run.font.name = Colors.FONT_EN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

    def _setup_headers_footers(self, doc: Document, variables: Dict[str, str]):
        """Setup headers and footers"""
        # Note: In python-docx, headers/footers need section access
        pass

    def _process_content(self, doc: Document, content: str):
        """Process markdown content to Word - FIXED: processes all tables"""
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                doc.add_paragraph()
                i += 1
                continue

            # Table detection
            if line.startswith("|") and "|" in line[1:]:
                i = self._process_table(doc, lines, i)  # Returns next line index
                continue

            # Headings
            if line.startswith("### "):
                heading = doc.add_heading(line[4:], level=3)
                self._format_heading(heading, 3)
            elif line.startswith("## "):
                heading = doc.add_heading(line[3:], level=2)
                self._format_heading(heading, 2)
            elif line.startswith("# "):
                heading = doc.add_heading(line[2:], level=1)
                self._format_heading(heading, 1)

            # Checkbox
            elif line.startswith("- [ ]") or line.startswith("- [x]"):
                self._add_checklist_item(doc, line)

            # Bullet
            elif line.startswith("- ") or line.startswith("* "):
                self._add_bullet_item(doc, line)

            # Separator
            elif "=" in line and len(line) < 20:
                i += 1
                continue

            # Regular paragraph
            else:
                self._add_formatted_paragraph(doc, line)

            i += 1

    def _process_table(self, doc: Document, lines: List[str], start_idx: int) -> int:
        """Process markdown table with styling - FIXED: returns next line index"""
        table_started = False
        table = None
        row_idx = 0

        for i in range(start_idx, len(lines)):
            line = lines[i].strip()

            if not line:
                continue

            # Skip separator lines
            if re.match(r"^\|[\s\-:|]+\|$", line):
                continue

            # End of table
            if not line.startswith("|"):
                if table_started:
                    return i  # Return current index for next processing
                continue

            cells = [c.strip() for c in line.split("|") if c.strip()]

            if not cells:
                continue

            if not table_started:
                # Create new table
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                table_started = True

                # Header row
                header_cells = table.rows[0].cells
                for j, cell_text in enumerate(cells):
                    if j < len(header_cells):
                        self._set_header_cell_text(header_cells[j], cell_text)
                row_idx = 1
            else:
                # Data row
                if len(cells) == len(table.columns):
                    row = table.add_row()
                    for j, cell_text in enumerate(cells):
                        if j < len(row.cells):
                            cell = row.cells[j]
                            self._set_data_cell_text(cell, cell_text)
                            self._apply_inline_formatting(cell, cell_text)

                    # Alternate row colors
                    if row_idx % 2 == 0:
                        for cell in row.cells:
                            shading = OxmlElement("w:shd")
                            shading.set(qn("w:fill"), "F1F5F9")  # Light slate
                            cell._tc.get_or_add_tcPr().append(shading)

                    row_idx += 1

        return len(lines)  # Return end of lines

    def _apply_inline_formatting(self, cell: _Cell, text: str):
        """Apply inline formatting based on markers"""
        para = cell.paragraphs[0]

        # Priority markers
        if "[必须]" in text or "[Must]" in text:
            for run in para.runs:
                run.font.color.rgb = Colors.PRIORITY_MUST
                run.font.bold = True
        elif "[应该]" in text or "[Should]" in text:
            for run in para.runs:
                run.font.color.rgb = Colors.PRIORITY_SHOULD
                run.font.bold = True
        elif "[可以]" in text or "[Could]" in text:
            for run in para.runs:
                run.font.color.rgb = Colors.PRIORITY_COULD

        # Status markers
        if "[Pass]" in text or "Pass" in text:
            for run in para.runs:
                if "Pass" in run.text:
                    run.font.color.rgb = Colors.STATUS_PASS
        elif "[Fail]" in text or "Fail" in text:
            for run in para.runs:
                if "Fail" in run.text:
                    run.font.color.rgb = Colors.STATUS_FAIL

    def _format_heading(self, heading, level: int):
        """Format heading with style"""
        for run in heading.runs:
            run.font.bold = True
            run.font.color.rgb = Colors.PRIMARY_DARK
            run.font.name = Colors.FONT_EN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)

    def _add_formatted_paragraph(self, doc: Document, text: str):
        """Add paragraph with inline formatting"""
        para = doc.add_paragraph()

        # Check for status markers
        if "[Pass]" in text:
            run = para.add_run(text.replace("[Pass]", "✓ "))
            run.font.color.rgb = Colors.STATUS_PASS
        elif "[Fail]" in text:
            run = para.add_run(text.replace("[Fail]", "✗ "))
            run.font.color.rgb = Colors.STATUS_FAIL
        elif "[必须]" in text:
            run = para.add_run(text)
            run.font.color.rgb = Colors.PRIORITY_MUST
            run.font.bold = True
        elif "[应该]" in text:
            run = para.add_run(text)
            run.font.color.rgb = Colors.PRIORITY_SHOULD
            run.font.bold = True
        else:
            run = para.add_run(text)
            run.font.color.rgb = Colors.TEXT_PRIMARY

        run.font.name = Colors.FONT_EN
        run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

    def _add_bullet_item(self, doc: Document, line: str):
        """Add bullet list item"""
        text = line[2:].strip()
        para = doc.add_paragraph(text, style="List Bullet")
        for run in para.runs:
            run.font.color.rgb = Colors.TEXT_PRIMARY
            run.font.name = Colors.FONT_EN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

    def _add_checklist_item(self, doc: Document, line: str):
        """Add checklist item"""
        is_checked = "[x]" in line[:4]
        text = line[4:].strip()
        para = doc.add_paragraph()

        if is_checked:
            run = para.add_run("☑ ")
            run.font.size = Pt(12)
        else:
            run = para.add_run("☐ ")
            run.font.size = Pt(12)

        text_run = para.add_run(text)
        text_run.font.color.rgb = Colors.TEXT_PRIMARY
        text_run.font.name = Colors.FONT_EN
        text_run._element.rPr.rFonts.set(qn("w:eastAsia"), Colors.FONT_CN)

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize filename"""
        if not name or name == "[Project Name]":
            return "Document"
        return re.sub(r'[<>:"/\\|?*]', "_", name)
